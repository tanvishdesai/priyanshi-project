"""
Generate a Word document (.docx) project report for
Computer Vision Lip Reading 2.0 — following the StrideX reference format.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── helpers ─────────────────────────────────────────────────────────────
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# heading styles
for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        hs.font.size = Pt(18)
    elif i == 2:
        hs.font.size = Pt(15)
    elif i == 3:
        hs.font.size = Pt(13)
    else:
        hs.font.size = Pt(12)


def add_centered(text, size=14, bold=True, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_normal(text, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    r = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        pass
    else:
        r.text = text
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
        set_cell_shading(cell, 'D9E2F3')
    # data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    return table


def page_break():
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ════════════════════════════════════════════════════════════════════════
add_centered('"Computer Vision Lip Reading System Using 3D Convolutional Neural Networks"', size=20, bold=True, space_after=24)
add_centered('A PROJECT REPORT', size=16, bold=True, space_after=12)
add_centered('Submitted by', size=13, bold=False, space_after=10)
add_centered('[YOUR NAME 1] [[ENROLLMENT_ID_1]]', size=13, bold=False, space_after=4)
add_centered('[YOUR NAME 2] [[ENROLLMENT_ID_2]]', size=13, bold=False, space_after=4)
add_centered('[YOUR NAME 3] [[ENROLLMENT_ID_3]]', size=13, bold=False, space_after=16)
add_centered('In fulfillment for the award of the degree\nof', size=13, bold=False, space_after=4)
add_centered('BACHELOR OF ENGINEERING', size=14, bold=True, space_after=4)
add_centered('in', size=13, bold=False, space_after=4)
add_centered('INFORMATION TECHNOLOGY', size=14, bold=True, space_after=16)
add_centered('LDRP Institute of Technology and Research, Gandhinagar\nKadi Sarva Vishwavidyalaya', size=13, bold=False, space_after=4)
add_centered('April-2025-26', size=13, bold=True, space_after=6)

page_break()

# ════════════════════════════════════════════════════════════════════════
#  CERTIFICATES  (3 copies, one per student)
# ════════════════════════════════════════════════════════════════════════
for idx in range(1, 4):
    add_centered('LDRP INSTITUTE OF TECHNOLOGY AND RESEARCH', size=14, bold=True, space_after=2)
    add_centered('GANDHINAGAR', size=13, bold=True, space_after=2)
    add_centered('IT Department', size=13, bold=False, space_after=12)
    add_centered('CERTIFICATE', size=16, bold=True, space_after=12)
    add_normal(
        f'This is to certify that the Project Work entitled "Computer Vision Lip Reading System Using '
        f'3D Convolutional Neural Networks" has been carried out by [YOUR NAME {idx}] '
        f'([ENROLLMENT_ID_{idx}]) under my guidance in fulfilment of the degree of Bachelor of Engineering '
        f'in Information Technology Semester-6 of Kadi Sarva Vishwavidyalaya during the academic year 2025-2026.'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    r1 = p.add_run('[Prof. Guide Name]')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run('                                        ')
    r3 = p.add_run('[Dr. HOD Name]')
    r3.bold = True
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)

    p2 = doc.add_paragraph()
    r4 = p2.add_run('Internal Guide')
    r4.font.name = 'Times New Roman'
    r4.font.size = Pt(12)
    r5 = p2.add_run('                                                      ')
    r6 = p2.add_run('Head of the Department')
    r6.font.name = 'Times New Roman'
    r6.font.size = Pt(12)

    p3 = doc.add_paragraph()
    r7 = p3.add_run('LDRP ITR')
    r7.font.name = 'Times New Roman'
    r7.font.size = Pt(12)
    r8 = p3.add_run('                                                              ')
    r9 = p3.add_run('LDRP ITR')
    r9.font.name = 'Times New Roman'
    r9.font.size = Pt(12)

    page_break()

# ════════════════════════════════════════════════════════════════════════
#  PRESENTATION PAGES
# ════════════════════════════════════════════════════════════════════════
for pnum in ['I', 'II']:
    add_centered(f'Presentation-{pnum} for Project-I', size=16, bold=True, space_after=20)
    add_normal('1.  Name & Signature of Internal Guide')
    add_normal('2.  Comments from Panel Members')
    add_normal('3.  Name & Signature of Panel Members')
    page_break()

# ════════════════════════════════════════════════════════════════════════
#  ACKNOWLEDGEMENT
# ════════════════════════════════════════════════════════════════════════
add_centered('ACKNOWLEDGEMENT', size=16, bold=True, space_after=16)
add_normal(
    'We would like to express our sincere gratitude to all those who have supported us in the '
    'successful completion of the project "Computer Vision Lip Reading System Using 3D Convolutional '
    'Neural Networks".'
)
add_normal(
    'First and foremost, we extend our heartfelt thanks to our project guide [PROF. GUIDE NAME] '
    '& Head of the department [DR. HOD NAME] whose valuable guidance, encouragement, and constructive '
    'suggestions have been a constant source of inspiration throughout this work.'
)
add_normal(
    'We are also thankful to the IT Department, LDRP-ITR for providing us with the facilities and '
    'resources required to carry out this project.'
)
add_normal(
    'Our deep appreciation goes to our classmates, friends, and family for their continuous '
    'motivation and support during this journey.'
)
add_normal(
    'Finally, we acknowledge with gratitude all the researchers and developers whose contributions '
    'in the fields of Computer Vision, Deep Learning, Speech Recognition, and Facial Landmark Detection '
    'have provided us with insights and references to complete this project successfully.'
)
add_normal(
    'This project has been a great learning experience, and we are truly grateful for the opportunity '
    'to explore, design, and implement this Computer Vision Lip Reading System.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('[Your Name 1] ([Enrollment_ID_1])\n[Your Name 2] ([Enrollment_ID_2])\n[Your Name 3] ([Enrollment_ID_3])')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

page_break()

# ════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ════════════════════════════════════════════════════════════════════════
add_centered('ABSTRACT', size=16, bold=True, space_after=16)

add_normal(
    'This project presents a Computer Vision-based Lip Reading System that leverages 3D Convolutional '
    'Neural Networks (3D CNNs) to recognize spoken words by analyzing lip movements captured through '
    'video. Unlike traditional speech recognition systems that rely on audio signals, this system '
    'operates entirely on visual input, making it applicable in noisy environments, for hearing-impaired '
    'individuals, and in scenarios where audio-based recognition is impractical.'
)
add_normal(
    'The system is trained on a custom-built dataset consisting of approximately 700 video clips '
    'totaling around 3 GB of data, covering 13 predefined English words: "here", "is", "a", "demo", '
    '"can", "you", "read", "my", "lips", "cat", "dog", "hello", and "bye". The dataset was manually '
    'collected using a webcam-based data collection pipeline that detects when a person is speaking by '
    'measuring the distance between upper and lower lip landmarks using a pre-trained facial landmark detector. '
    'Each captured word is represented as a sequence of 22 preprocessed lip-region frames, each sized '
    '80x112 pixels in RGB color space.'
)
add_normal(
    'The core model architecture is a 3D Convolutional Neural Network built using TensorFlow and Keras, '
    'featuring three 3D convolutional layers (with 8, 32, and 256 filters respectively), 3D max-pooling layers, '
    'L2 regularization, and multiple fully connected layers with dropout for regularization. '
    'The model was trained using the Adam optimizer with categorical cross-entropy loss over 20 epochs, '
    'achieving a training accuracy of 95.7% and a validation accuracy of 98.5%, demonstrating strong '
    'classification performance across all 13 word classes.'
)
add_normal(
    'The system includes a real-time live prediction mode where a webcam feed is processed frame-by-frame '
    'to detect speaking activity, capture lip frames, and feed them through the trained model for instant '
    'word prediction. Additionally, a file-based prediction mode allows inference on pre-recorded video files. '
    'The preprocessing pipeline includes LAB color space conversion, CLAHE contrast enhancement, Gaussian blurring, '
    'bilateral filtering, and sharpening to normalize lip appearance across varying lighting conditions.'
)
add_normal(
    'By combining computer vision techniques with deep learning, this project demonstrates that visual-only '
    'lip reading is a viable approach for word-level speech recognition, with potential applications in '
    'assistive technology, security surveillance, and human-computer interaction.'
)

page_break()

# ════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════
add_centered('TABLE OF CONTENTS', size=16, bold=True, space_after=16)

toc_data = [
    ['1', 'INTRODUCTION', ''],
    ['', '1.1 Introduction', ''],
    ['', '1.2 Aims and Objective of the Work', ''],
    ['', '    1.2.1 Aims', ''],
    ['', '    1.2.2 Objectives', ''],
    ['', '1.3 Brief Literature Review', ''],
    ['', '    1.3.1 Visual Speech Recognition', ''],
    ['', '    1.3.2 Deep Learning for Lip Reading', ''],
    ['', '    1.3.3 Facial Landmark Detection', ''],
    ['', '1.4 Problem Definition', ''],
    ['', '1.5 Plan of the Work', ''],
    ['2', 'TECHNOLOGY AND LITERATURE REVIEW', ''],
    ['', '2.1 Technologies', ''],
    ['', '    2.1.1 Programming Language & Runtime', ''],
    ['', '    2.1.2 Deep Learning Framework', ''],
    ['', '    2.1.3 Computer Vision Libraries', ''],
    ['', '    2.1.4 Data Processing & Utilities', ''],
    ['', '2.2 Literature Review', ''],
    ['', '    2.2.1 Contextual Background', ''],
    ['', '    2.2.2 Existing Solutions & Market Analysis', ''],
    ['', '2.3 Problem Statement & Gaps', ''],
    ['', '2.4 The Role of Deep Learning Intervention (Project\'s Contribution)', ''],
    ['3', 'SYSTEM REQUIREMENTS STUDY', ''],
    ['', '3.1 User Characteristics', ''],
    ['', '    3.1.1 Hearing-Impaired Individuals (Primary Users)', ''],
    ['', '    3.1.2 Researchers & Developers (Secondary Users)', ''],
    ['', '3.2 Hardware and Software Requirements', ''],
    ['', '    3.2.1 Hardware Requirements', ''],
    ['', '    3.2.2 Software Requirements', ''],
    ['', '3.3 Assumptions and Dependencies', ''],
    ['', '    3.3.1 Assumptions', ''],
    ['', '    3.3.2 Dependencies', ''],
    ['4', 'SYSTEM DIAGRAMS', ''],
    ['', '4.1 Model Architecture Diagram', ''],
    ['', '4.2 Data Pipeline Flowchart', ''],
    ['', '4.3 Use Case Diagram', ''],
    ['', '4.4 Sequence Diagram', ''],
    ['', '4.5 Activity Diagram', ''],
    ['', '4.6 Data Flow Diagram', ''],
    ['', '    4.6.1 Context-Level-0 DFD', ''],
    ['', '    4.6.2 Context-Level-1 DFD', ''],
    ['', '    4.6.3 Context-Level-2 DFD For Data Collection', ''],
    ['', '    4.6.4 Context-Level-2 DFD For Prediction', ''],
    ['5', 'DATA DICTIONARY', ''],
    ['', '5.1 Overview', ''],
    ['', '5.2 Data Schema', ''],
    ['', '    5.2.1 Collected Data Structure', ''],
    ['', '    5.2.2 Model Input Tensor', ''],
    ['', '5.3 Data Integrity and Processing', ''],
    ['', '5.4 Dataset Description and Statistics', ''],
    ['', '    5.4.1 Dataset Overview', ''],
    ['', '    5.4.2 Per-Class Sample Distribution', ''],
    ['', '    5.4.3 Dataset Characteristics', ''],
    ['', '    5.4.4 Data Collection Methodology', ''],
    ['6', 'RESULT, DISCUSSION AND CONCLUSION', ''],
    ['', '6.1 Result', ''],
    ['', '6.2 Conclusion', ''],
    ['7', 'BIBLIOGRAPHY', ''],
    ['', '7.1 References', ''],
    ['', '    7.1.1 Academic Papers & Journals', ''],
    ['', '    7.1.2 Technical Documentation', ''],
    ['', '    7.1.3 Web Resources & Libraries', ''],
    ['', '    7.1.4 Similar Systems (Case Studies)', ''],
]
add_table(['NO', 'CHAPTER NAME', 'PAGE NO'], toc_data)

page_break()

# ════════════════════════════════════════════════════════════════════════
#  LIST OF FIGURES
# ════════════════════════════════════════════════════════════════════════
add_centered('LIST OF FIGURES', size=16, bold=True, space_after=16)
fig_data = [
    ['1', 'Model Architecture Diagram (3D CNN)', ''],
    ['2', 'Data Pipeline Flowchart', ''],
    ['3', 'Use Case Diagram', ''],
    ['4', 'Sequence Diagram (Live Prediction Flow)', ''],
    ['5', 'Activity Diagram (Data Collection & Prediction Workflow)', ''],
    ['6', 'Context-Level-0 DFD', ''],
    ['7', 'Context-Level-1 DFD', ''],
    ['8', 'Context-Level-2 DFD For Data Collection', ''],
    ['9', 'Context-Level-2 DFD For Prediction', ''],
    ['10', 'Training & Validation Accuracy/Loss Graph', ''],
    ['11', 'Confusion Matrix Heatmap', ''],
    ['12', 'ROC-AUC Curve', ''],
    ['13', 'Classification Metrics Table', ''],
    ['14', 'App Screenshot – Data Collection Mode', ''],
    ['15', 'App Screenshot – Live Prediction Mode', ''],
    ['16', 'Sample Lip Frames for Different Words', ''],
]
add_table(['NO', 'NAME', 'PAGE NO'], fig_data)

page_break()

# ════════════════════════════════════════════════════════════════════════
#  LIST OF TABLES
# ════════════════════════════════════════════════════════════════════════
add_centered('LIST OF TABLES', size=16, bold=True, space_after=16)
tbl_data = [
    ['1', 'Collected Data Structure', ''],
    ['2', 'Model Input Tensor Schema', ''],
    ['3', 'Label Encoding Mapping', ''],
    ['4', 'Per-Class Classification Metrics', ''],
    ['5', 'Per-Class Sample Distribution', ''],
    ['6', 'Dataset Characteristics Summary', ''],
]
add_table(['NO', 'NAME', 'PAGE NO'], tbl_data)

page_break()

# ════════════════════════════════════════════════════════════════════════
#  CHAPTER 1: INTRODUCTION
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('1    INTRODUCTION', level=1)
doc.add_heading('1.1  INTRODUCTION', level=2)
doc.add_heading('1.2  AIMS AND OBJECTIVE OF THE WORK', level=2)
doc.add_heading('1.3  BRIEF LITERATURE REVIEW', level=2)
doc.add_heading('1.4  PROBLEM DEFINITION', level=2)
doc.add_heading('1.5  PLAN OF THE WORK', level=2)

page_break()

# ── 1.1 Introduction ──
doc.add_heading('1.1. Introduction', level=2)
add_normal(
    'Lip reading, also known as visual speech recognition (VSR), is the process of interpreting '
    'spoken language by visually analyzing the movements of the lips, face, and tongue. It is a '
    'skill naturally employed by hearing-impaired individuals and has gained significant research '
    'attention in the fields of computer vision and artificial intelligence. This project presents '
    'a Computer Vision Lip Reading System that uses 3D Convolutional Neural Networks (3D CNNs) to '
    'automatically recognize spoken words from video sequences of lip movements.'
)
add_normal(
    'The system operates entirely on visual input—no audio signal is required. A webcam captures '
    'the user\'s face in real-time, and a pre-trained facial landmark detector identifies 68 key facial '
    'points, from which the lip region is extracted. The system monitors the '
    'distance between the upper and lower lip landmarks to detect when the '
    'user is speaking. When speech is detected, the system captures a sequence of 22 consecutive '
    'lip-region frames, each preprocessed to a standardized 80×112 pixel RGB image with contrast '
    'enhancement and noise reduction.'
)
add_normal(
    'These 22-frame sequences form the input to a 3D Convolutional Neural Network that has been '
    'trained to classify lip movement patterns into one of 13 predefined English words: "here", '
    '"is", "a", "demo", "can", "you", "read", "my", "lips", "cat", "dog", "hello", and "bye". '
    'The 3D CNN architecture is specifically chosen because it can capture both spatial features '
    '(the shape and position of the lips in each frame) and temporal features (how the lip shape '
    'changes across the sequence of frames), making it ideal for video-based classification tasks.'
)
add_normal(
    'The model was trained on a custom-built dataset of approximately 700 video clips (~3 GB of data), '
    'collected manually using a purpose-built data collection module. The training process achieved '
    'a 95.7% training accuracy and 98.5% validation accuracy, demonstrating strong classification '
    'performance. The system includes both a live prediction mode (real-time webcam inference) and '
    'a file-based prediction mode (inference on pre-recorded videos), making it versatile for both '
    'interactive and batch processing scenarios.'
)
add_normal(
    'By leveraging advances in facial landmark detection, image preprocessing, and 3D convolutional '
    'neural networks, this project demonstrates that visual-only lip reading is a viable and accurate '
    'approach for word-level speech recognition, with applications in assistive technology for the '
    'hearing impaired, security and surveillance systems, noisy environment communication, and '
    'human-computer interaction research.'
)

page_break()

# ── 1.2 Aims and Objectives ──
doc.add_heading('1.2. Aims and Objective of the Work', level=2)
doc.add_heading('1.2.1. Aims of "Computer Vision Lip Reading System"', level=3)
add_normal(
    'The Computer Vision Lip Reading System is designed to recognize spoken words from visual lip '
    'movement patterns using deep learning. Based on the project\'s design and implementation, its '
    'primary aims are:'
)
add_bullet('Visual Speech Recognition: To develop a system that can accurately recognize spoken words '
           'solely from visual analysis of lip movements, without requiring any audio input.')
add_bullet('Automated Lip Detection & Frame Capture: To build an intelligent pipeline that automatically '
           'detects when a person is speaking by analyzing facial landmark distances, captures the relevant '
           'lip-region frames, and preprocesses them for classification.')
add_bullet('Deep Learning Classification: To train a 3D Convolutional Neural Network capable of learning '
           'spatiotemporal features from lip movement sequences and classifying them into 13 predefined '
           'word categories with high accuracy.')
add_bullet('Real-Time Inference: To enable real-time word prediction from a live webcam feed, providing '
           'immediate visual feedback of the recognized word on screen.')
add_normal(
    'In short, the project aims to demonstrate that computer vision and deep learning can be combined '
    'to build a practical, visual-only lip reading system that works in real-time.'
)

doc.add_heading('1.2.2. Objectives of "Computer Vision Lip Reading System"', level=3)
add_normal(
    'The specific technical objectives of this project are:'
)
add_bullet('Implement Facial Landmark Detection: Utilize a pre-trained shape predictor '
           'with 68 facial landmarks to detect faces and extract the lip region '
           'from each video frame with sub-pixel accuracy.')
add_bullet('Build a Custom Dataset: Design and implement a data collection pipeline that '
           'uses lip-distance thresholds to detect speech onset/offset, captures exactly 22 frames per word '
           '(including pre-speech buffer frames), and saves them as structured data for training.')
add_bullet('Develop Image Preprocessing Pipeline: Implement a robust preprocessing chain including LAB '
           'color space conversion, CLAHE (Contrast Limited Adaptive Histogram Equalization) for contrast '
           'enhancement, Gaussian blurring, bilateral filtering, and sharpening to normalize lip appearance '
           'across varying lighting conditions.')
add_bullet('Design and Train a 3D CNN Model: Architect and train a sequential deep learning model with three '
           '3D convolutional layers (8, 32, 256 filters), 3D max-pooling layers, L2 regularization, fully connected layers, '
           'dropout regularization, and softmax output for 13-class classification, achieving >95% validation accuracy.')
add_bullet('Implement Live Prediction System: Build a real-time prediction module that '
           'captures webcam frames, detects speech activity, accumulates 22 preprocessed lip frames, feeds '
           'them through the trained model, and displays the predicted word on screen.')
add_bullet('Implement File-Based Prediction: Build a video file prediction module that '
           'processes pre-recorded videos frame-by-frame using the same pipeline and saves annotated output '
           'videos with predicted word overlays.')

page_break()

# ── 1.3 Literature Review ──
doc.add_heading('1.3. Brief Literature Review', level=2)
add_normal(
    'The development of this Lip Reading System is grounded in existing research and innovations '
    'across the fields of visual speech recognition, deep learning for video analysis, and facial '
    'landmark detection.'
)

doc.add_heading('1.3.1. Visual Speech Recognition', level=3)
add_normal(
    'Visual speech recognition (VSR), also known as automatic lip reading, has been an active area '
    'of research since the 1980s. Early systems relied on hand-crafted features such as lip contour '
    'shapes and optical flow vectors. Potamianos et al. (2003) provided a comprehensive survey of '
    'audio-visual speech recognition, establishing that visual information from lip movements can '
    'significantly improve speech recognition accuracy, especially in noisy environments. The GRID '
    'corpus (Cooke et al., 2006) became a standard benchmark dataset, containing 34,000 short video '
    'clips of 34 speakers, enabling systematic evaluation of VSR systems.'
)

doc.add_heading('1.3.2. Deep Learning for Lip Reading', level=3)
add_normal(
    'The breakthrough application of deep learning to lip reading came with the LipNet model '
    '(Assael et al., 2016), which achieved 95.2% accuracy on the GRID corpus using an end-to-end '
    'deep learning architecture combining Spatiotemporal CNNs, GRUs, and CTC loss. This was followed '
    'by "Lip Reading Sentences in the Wild" (Chung et al., 2017) from DeepMind, which demonstrated '
    'that deep learning models could surpass professional human lip readers. The use of 3D Convolutional '
    'Neural Networks (3D CNNs) for video analysis was pioneered by Tran et al. (2015) with the C3D model, '
    'which showed that 3D convolutions can effectively learn spatiotemporal features from video sequences—'
    'the same architectural principle employed in this project.'
)

doc.add_heading('1.3.3. Facial Landmark Detection', level=3)
add_normal(
    'The facial landmark detector used in this project, based on the work of Kazemi and Sullivan (2014) '
    '"One Millisecond Face Alignment with an Ensemble of Regression Trees," provides real-time '
    'detection of 68 facial landmarks with high accuracy. This detector is widely used in face '
    'recognition, emotion detection, and lip reading applications. The 68-landmark model identifies '
    'key points around the jaw, eyebrows, nose, eyes, and mouth—with the mouth-region landmarks specifically '
    'representing the outer and inner lip contours, which are critical for extracting the lip region '
    'of interest (ROI) used in this project.'
)

page_break()

# ── 1.4 Problem Definition ──
doc.add_heading('1.4. Problem Definition', level=2)
add_normal(
    'The current landscape of speech recognition and accessibility technology faces several '
    'critical challenges that this project seeks to address:'
)
add_bullet('Audio Dependency of Speech Recognition: Mainstream speech recognition systems (Google Speech-to-Text, '
           'Apple Siri, Amazon Alexa) rely entirely on audio input. They fail in noisy environments (factories, '
           'concerts, busy streets), underwater, behind glass barriers, or when the speaker is muted. There is '
           'no widely available consumer system that can recognize speech from visual input alone.')
add_bullet('Limited Accessibility for Hearing-Impaired: While sign language recognition has received research '
           'attention, automated lip reading—which would allow hearing-impaired individuals to understand speakers '
           'who do not know sign language—remains an underexplored application area with no mainstream consumer solution.')
add_bullet('Dataset Scarcity for Word-Level Lip Reading: Most existing lip reading datasets (GRID, LRW, LRS) '
           'are designed for sentence-level or phoneme-level recognition, often with controlled lighting and '
           'studio conditions. There is a scarcity of simple, word-level datasets suitable for training lightweight '
           'models that can run in real-time on consumer hardware.')
add_bullet('Real-Time Processing Challenge: Performing facial landmark detection, lip region extraction, image '
           'preprocessing, and deep learning inference in real-time from a webcam feed presents computational '
           'challenges. The pipeline must be efficient enough to process frames at video rate (15-30 FPS) while '
           'maintaining classification accuracy.')

page_break()

# ── 1.5 Plan of the Work ──
doc.add_heading('1.5. Plan of the Work', level=2)
add_normal(
    'To ensure the successful development of the Computer Vision Lip Reading System, the project '
    'work is divided into the following strategic phases:'
)

add_normal('Phase 1: Requirement Analysis & Feasibility', bold=True)
add_bullet('Reviewing literature on visual speech recognition, 3D CNNs, and facial landmark detection.')
add_bullet('Defining the vocabulary: 13 common English words.')
add_bullet('Determining the technical stack: Python, TensorFlow/Keras, OpenCV, dlib, NumPy.')
add_bullet('Evaluating feasibility of real-time lip reading on consumer hardware.')

add_normal('Phase 2: Data Collection & Dataset Creation', bold=True)
add_bullet('Designing the data collection pipeline with automatic speech detection via lip-distance thresholds.')
add_bullet('Implementing the data collection module with calibration mode, circular frame buffers, and standardized 22-frame output.')
add_bullet('Collecting ~700 video clips across 13 word classes using webcam recordings.')
add_bullet('Organizing data into labeled folders with structured array files and video files.')

add_normal('Phase 3: Model Design & Training', bold=True)
add_bullet('Designing the 3D CNN architecture with 3D convolutional, max-pooling, fully connected, and dropout layers.')
add_bullet('Implementing data loading, label encoding, and train/test split (80/20).')
add_bullet('Training the model for 20 epochs with Adam optimizer and categorical cross-entropy loss.')
add_bullet('Evaluating with confusion matrix, classification report, ROC-AUC curves, and balanced accuracy.')

add_normal('Phase 4: Inference System Development', bold=True)
add_bullet('Building the real-time prediction module for webcam-based word prediction.')
add_bullet('Building the file-based prediction module for batch inference on pre-recorded video files.')
add_bullet('Implementing the same preprocessing pipeline used during training for inference consistency.')

add_normal('Phase 5: Testing & Validation', bold=True)
add_bullet('Testing with various speakers, lighting conditions, and camera angles.')
add_bullet('Validating model accuracy against the test set (98.5% validation accuracy achieved).')
add_bullet('Recording demonstration videos to showcase system capabilities.')

page_break()

# ════════════════════════════════════════════════════════════════════════
#  CHAPTER 2: TECHNOLOGY AND LITERATURE REVIEW
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('2    TECHNOLOGY AND LITERATURE REVIEW', level=1)
doc.add_heading('2.1  TECHNOLOGIES', level=2)
doc.add_heading('2.2  LITERATURE REVIEW', level=2)

page_break()

# ── 2.1 Technologies ──
doc.add_heading('2.1. Technologies', level=2)

doc.add_heading('2.1.1. Programming Language & Runtime', level=3)
add_normal(
    'Python 3 (v3.11/3.12): The primary programming language for the entire project. Python is '
    'the de facto standard for machine learning and computer vision projects due to its extensive '
    'ecosystem of scientific computing libraries, readable syntax, and strong community support. '
    'All data collection, preprocessing, model training, and inference scripts are written in Python.',
)

doc.add_heading('2.1.2. Deep Learning Framework', level=3)
add_normal(
    'TensorFlow (v2.12.0): Google\'s open-source machine learning framework used as the backend '
    'for building, training, and running the 3D CNN model. TensorFlow provides GPU acceleration, '
    'automatic differentiation, and production-ready model serving capabilities.'
)
add_normal(
    'Keras (v2.12.0): The high-level neural network API running on top of TensorFlow. Keras provides '
    'the Sequential model API, layer primitives (Conv3D, MaxPooling3D, Dense, Dropout, Flatten), '
    'optimizers (Adam), loss functions (categorical_crossentropy), and training utilities (model.fit, '
    'model.predict) used throughout the project. Key Keras components used include:'
)
add_bullet('Conv3D: 3D convolutional layers that apply filters across both spatial (height, width) and '
           'temporal (frame sequence) dimensions, enabling the model to learn spatiotemporal features '
           'from video sequences.')
add_bullet('MaxPooling3D: Downsampling layers that reduce spatial and temporal dimensions, decreasing '
           'computational cost while retaining the most important features.')
add_bullet('Dense: Fully connected layers for classification. The final Dense layer uses softmax activation '
           'to output probability distributions across the 13 word classes.')
add_bullet('Dropout: Regularization layers (rate=0.5) that randomly zero out neurons during training to '
           'prevent overfitting.')
add_bullet('L2 Regularization (kernel_regularizer): Applied to Conv3D layers with a factor of 0.001 to '
           'penalize large weights and improve generalization.')

doc.add_heading('2.1.3. Computer Vision Libraries', level=3)
add_normal(
    'OpenCV (v4.6.0.66): The primary computer vision library used for webcam capture (VideoCapture), '
    'image color space conversion (BGR to Grayscale, BGR to LAB), image preprocessing (CLAHE, '
    'GaussianBlur, bilateralFilter, filter2D for sharpening), image resizing, drawing annotations '
    '(circles on landmarks, text overlays), and video writing (VideoWriter for output files).'
)
add_normal(
    'dlib (v19.24.1): A C++ toolkit with Python bindings used for two critical functions:'
)
add_bullet('Face Detection: dlib.get_frontal_face_detector() — a HOG+SVM-based frontal face detector '
           'that identifies face bounding boxes in grayscale images.')
add_bullet('Facial Landmark Prediction: dlib.shape_predictor() with the pre-trained "shape_predictor_68_face_landmarks" '
           'model (face_weights.dat) — predicts 68 facial landmark coordinates, from which lip landmarks '
           '(48-61) are extracted for lip region cropping.')
add_normal(
    'Pillow (PIL) (v9.4.0): Python Imaging Library used for creating RGB images from numpy pixel arrays '
    'during the data saving process in the collection pipeline.'
)

doc.add_heading('2.1.4. Data Processing & Utilities', level=3)
add_normal(
    'NumPy (v1.23.5): The fundamental numerical computing library used for array operations throughout '
    'the project — storing frame sequences as 4D arrays (frames × height × width × channels), '
    'preprocessing kernels (sharpening filter), and model input/output processing.'
)
add_normal(
    'scikit-learn: Used for data splitting (train_test_split with 80/20 ratio), label encoding '
    '(LabelEncoder), evaluation metrics (classification_report, balanced_accuracy_score, confusion_matrix, '
    'roc_curve, auc), and label binarization for ROC analysis.'
)
add_normal(
    'imageio (v2.27.0): Used for reading image files and writing MP4 video files from frame sequences '
    'during the data collection process (imageio.mimsave).'
)
add_normal(
    'matplotlib: Used in the training notebook for plotting training/validation accuracy and loss curves, '
    'class distribution histograms, and ROC-AUC curves.'
)
add_normal(
    'seaborn: Used for generating the confusion matrix heatmap visualization in the training notebook.'
)

page_break()

# ── 2.2 Literature Review ──
doc.add_heading('2.2. Literature Review', level=2)

doc.add_heading('2.2.1. Contextual Background', level=3)
add_normal(
    'The Communication Barrier for the Hearing Impaired: The World Health Organization (WHO) reports '
    'that over 430 million people worldwide suffer from disabling hearing loss, a number projected to '
    'rise to 700 million by 2050. While hearing aids and cochlear implants address some cases, many '
    'individuals rely on lip reading as a primary communication method. However, human lip reading is '
    'extremely difficult — studies show that even professional lip readers achieve only 40-60% accuracy '
    'on unrestricted vocabulary. An automated lip reading system could significantly enhance communication '
    'accessibility.'
)
add_normal(
    'The Rise of Deep Learning in Visual Recognition: The deep learning revolution, beginning with '
    'AlexNet (Krizhevsky et al., 2012) winning ImageNet, has transformed computer vision. Convolutional '
    'Neural Networks (CNNs) have become the standard for image classification, and their extension to '
    'video analysis via 3D CNNs (Tran et al., 2015) has enabled temporal feature learning from video '
    'sequences. This project leverages these advances to apply 3D CNNs to the specific problem of lip '
    'movement classification.'
)
add_normal(
    'Audio-Visual Speech Recognition: The McGurk Effect (McGurk & MacDonald, 1976) demonstrated that '
    'visual information from lip movements significantly influences speech perception. This finding has '
    'driven decades of research into audio-visual speech recognition (AVSR), where visual lip information '
    'is fused with audio signals to improve recognition accuracy in noisy conditions (Potamianos et al., 2003).'
)

doc.add_heading('2.2.2. Existing Solutions & Market Analysis', level=3)
add_normal(
    'LipNet (Assael et al., 2016): The first end-to-end deep learning model for sentence-level lip reading, '
    'achieving 95.2% accuracy on the GRID corpus. LipNet uses Spatiotemporal CNNs + GRUs + CTC loss, but '
    'requires large-scale labeled video datasets and significant computational resources.'
)
add_normal(
    'DeepMind\'s Visual Speech Recognition (Chung et al., 2017): "Lip Reading Sentences in the Wild" '
    'demonstrated performance surpassing professional human lip readers on the LRS dataset. However, the '
    'system requires massive training data (100,000+ utterances) and is designed for sentence-level recognition, '
    'making it impractical for lightweight, word-level applications.'
)
add_normal(
    'Google Speech-to-Text, Apple Siri, Amazon Alexa: These mainstream speech recognition systems operate '
    'entirely on audio input and cannot function in visually-observed-only scenarios. They offer no lip '
    'reading capability.'
)
add_normal(
    'Commercial Lip Reading Systems: Companies such as Liopa (UK) have developed lip reading technology '
    'for healthcare settings, but these are proprietary, expensive, and designed for specific clinical use cases '
    'rather than general-purpose word recognition.'
)

page_break()

# ── 2.3 Problem Statement & Gaps ──
doc.add_heading('2.3. Problem Statement & Gaps', level=2)
add_normal(
    'Accessibility Gap: Despite 430 million people with hearing loss, no mainstream consumer application '
    'provides real-time visual lip reading to assist communication. Current assistive technology focuses '
    'on hearing aids and sign language translation, leaving lip reading automation as an unmet need.',
    bold=False
)
add_normal(
    'Complexity Gap: Existing deep learning lip reading systems (LipNet, DeepMind) require massive '
    'datasets, sentence-level architectures, and significant computational resources. There is a gap '
    'for lightweight, word-level models that can run in real-time on standard consumer hardware with '
    'custom-built, smaller datasets.'
)
add_normal(
    'Dataset Gap: Most lip reading datasets (GRID, LRW, LRS) feature controlled studio conditions and '
    'are designed for sentence or phoneme-level recognition. There is a scarcity of word-level datasets '
    'collected under natural webcam conditions suitable for training lightweight models.'
)
add_normal(
    'Audio-Free Gap: The vast majority of speech recognition systems require audio input. A purely visual '
    'system that operates without any microphone is needed for scenarios such as: communication through '
    'glass barriers, surveillance footage analysis, space environments, and any context where audio '
    'recording is impractical or prohibited.'
)

# ── 2.4 Project\'s Contribution ──
doc.add_heading('2.4. The Role of Deep Learning Intervention (The Project\'s Contribution)', level=2)
add_normal(
    'From Audio to Visual: This project shifts speech recognition from audio-dependent to purely visual, '
    'demonstrating that a well-designed 3D CNN trained on lip movement sequences can achieve high accuracy '
    '(98.5% validation) for word-level classification without any audio signal.'
)
add_normal(
    'Custom Dataset Creation: The project contributes a novel word-level lip reading dataset of ~700 '
    'video clips across 13 words, collected under natural webcam conditions with an automated collection '
    'pipeline. This dataset and methodology can be replicated and extended by other researchers.'
)
add_normal(
    'Lightweight Real-Time Architecture: Unlike heavyweight sentence-level models, this project demonstrates '
    'that a relatively simple 3D CNN architecture (3 Conv3D layers + Dense layers) can achieve excellent '
    'word-level classification performance while being light enough for real-time webcam inference.'
)
add_normal(
    'End-to-End Pipeline: The project delivers a complete pipeline from data collection → preprocessing → '
    'model training → real-time inference → file-based inference, serving as a practical reference '
    'implementation for visual speech recognition research.'
)
add_normal(
    'Robust Preprocessing: The LAB-CLAHE-Gaussian-Bilateral-Sharpening preprocessing pipeline addresses '
    'the real-world challenge of varying lighting conditions during webcam-based lip capture, contributing '
    'a practical solution for normalizing lip appearance in uncontrolled environments.'
)

page_break()

# ════════════════════════════════════════════════════════════════════════
#  CHAPTER 3: SYSTEM REQUIREMENTS STUDY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('3    SYSTEM REQUIREMENTS STUDY', level=1)
doc.add_heading('3.1  USER CHARACTERISTICS', level=2)
doc.add_heading('3.2  HARDWARE AND SOFTWARE REQUIREMENTS', level=2)
doc.add_heading('3.3  ASSUMPTIONS AND DEPENDENCIES', level=2)

page_break()

# ── 3.1 User Characteristics ──
doc.add_heading('3.1. User Characteristics', level=2)
add_normal(
    'The Computer Vision Lip Reading System is designed for two distinct user groups, each with '
    'specific roles and technical aptitudes.'
)

doc.add_heading('3.1.1. Hearing-Impaired Individuals & Communication Aid Users (Primary Users)', level=3)
add_normal('Role: End users who benefit from real-time visual lip reading for communication assistance.', bold=True)
add_normal('Characteristics:')
add_bullet('Demographics: Individuals with partial or complete hearing loss of all ages.')
add_bullet('Technical Expertise: Low to Medium. The system is designed for simple usage — run a script and speak into the webcam.')
add_bullet('Frequency of Use: Regular (daily conversations, meetings, educational settings).')
add_normal('Key Privileges:')
add_bullet('Use live prediction mode to see real-time word recognition from webcam.')
add_bullet('View predicted words displayed on screen.')
add_bullet('Reset prediction history for new conversations.')

doc.add_heading('3.1.2. Researchers & Developers (Secondary Users)', level=3)
add_normal('Role: Users who extend, modify, or build upon the system for research or development.', bold=True)
add_normal('Characteristics:')
add_bullet('Demographics: Computer science students, machine learning researchers, accessibility technology developers.')
add_bullet('Technical Expertise: High. Familiar with Python, deep learning, and computer vision.')
add_bullet('Frequency of Use: Periodic (during research, model retraining, dataset expansion).')
add_normal('Key Privileges:')
add_bullet('Use data collection mode to record new training data for additional words.')
add_bullet('Retrain the model with expanded datasets using the Jupyter notebook.')
add_bullet('Modify preprocessing parameters and model architecture.')
add_bullet('Use file-based prediction for batch processing of video files.')

page_break()

# ── 3.2 Hardware and Software Requirements ──
doc.add_heading('3.2. Hardware and Software Requirements', level=2)

doc.add_heading('3.2.1. Hardware Requirements', level=3)
add_normal('Client-Side (User Machine)', bold=True)
add_bullet('Processor: Intel Core i5 or equivalent (quad-core 2.0 GHz minimum). GPU recommended for training (NVIDIA CUDA-compatible).')
add_bullet('RAM: Minimum 8 GB (16 GB recommended for model training).')
add_bullet('Webcam: Built-in or external USB webcam with minimum 480p resolution (720p recommended).')
add_bullet('Storage: Minimum 5 GB free space (3 GB for dataset, 1 GB for model weights, 1 GB for dependencies).')
add_bullet('Display: Standard monitor for viewing the webcam feed and prediction results.')

add_normal('Training Environment (Optional)', bold=True)
add_bullet('The training was performed on Kaggle using cloud GPU (NVIDIA Tesla P100 or T4).')
add_bullet('Local training requires an NVIDIA GPU with at least 4 GB VRAM and CUDA toolkit installed.')

doc.add_heading('3.2.2. Software Requirements', level=3)
add_normal('Development/Runtime Environment', bold=True)
add_bullet('Operating System: Windows 10/11, macOS (Catalina or later), or Linux (Ubuntu 18.04+).')
add_bullet('Python: Version 3.11 or 3.12.')
add_bullet('Code Editor: Visual Studio Code, Cursor IDE, or Jupyter Notebook (for training).')

add_normal('Core Dependencies', bold=True)
add_table(
    ['Package', 'Version', 'Purpose'],
    [
        ['TensorFlow', '2.12.0', 'Deep learning framework for model building and inference'],
        ['Keras', '2.12.0', 'High-level neural network API (included in TensorFlow)'],
        ['OpenCV', '4.6.0.66', 'Computer vision: webcam, image processing, video I/O'],
        ['dlib', '19.24.1', 'Face detection and facial landmark prediction (68 points)'],
        ['NumPy', '1.23.5', 'Numerical array operations for frame data'],
        ['Pillow', '9.4.0', 'Image creation from numpy arrays during data saving'],
        ['imageio', '2.27.0', 'Video file creation from frame sequences'],
        ['scikit-learn', '—', 'Data splitting, label encoding, evaluation metrics'],
        ['matplotlib', '—', 'Training curve visualization'],
        ['seaborn', '—', 'Confusion matrix heatmap visualization'],
        ['h5py', '3.8.0', 'HDF5 file I/O for model weight storage (.h5 files)'],
    ]
)

page_break()

# ── 3.3 Assumptions and Dependencies ──
doc.add_heading('3.3. Assumptions and Dependencies', level=2)

doc.add_heading('3.3.1. Assumptions', level=3)
add_bullet('Webcam Availability: It is assumed that the user\'s machine has a functioning webcam accessible by the system. External USB webcams should work but may require driver installation.')
add_bullet('Frontal Face Orientation: The system assumes the user is facing the camera approximately frontally. Profile or extreme angle views will cause the facial landmark detector to fail, resulting in no lip detection.')
add_bullet('Adequate Lighting: It is assumed that the user\'s face is reasonably well-lit. While the contrast enhancement preprocessing helps normalize lighting, extremely dark or heavily backlit conditions may degrade accuracy.')
add_bullet('Single Speaker: The current implementation processes one face at a time (the first detected face). It is assumed that only one speaker is present in the camera frame.')
add_bullet('English Words Only: The system is trained exclusively on 13 English words. It assumes all spoken input belongs to this predefined vocabulary; out-of-vocabulary words will be misclassified into the closest known class.')
add_bullet('Consistent Speaking Style: The system assumes a natural speaking pace and mouth movement. Exaggerated or whispered speech may not match the training data distribution.')

doc.add_heading('3.3.2. Dependencies', level=3)
add_bullet('Pre-trained Facial Landmark Model: The system depends on a pre-trained 68-point facial landmark model file. This file must be present in the designated model directory. Without it, face detection and lip extraction fail entirely.')
add_bullet('Trained Neural Network Weights: The trained 3D CNN weight files must be present in the model directory. Without pre-trained weights, the system cannot perform prediction.')
add_bullet('Webcam Interface: The webcam capture depends on the computer vision library interfacing with the system\'s camera driver. Compatibility issues may arise on certain operating systems or with specific webcam hardware.')
add_bullet('Package Compatibility: The project depends on specific versions of the deep learning framework, facial landmark library, and computer vision library that must be mutually compatible.')
add_bullet('GPU Drivers (Training Only): If retraining the model locally, the system depends on NVIDIA CUDA drivers and cuDNN library being correctly installed for GPU-accelerated training.')

page_break()

# ════════════════════════════════════════════════════════════════════════
#  CHAPTER 4: SYSTEM DIAGRAMS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('4    SYSTEM DIAGRAMS', level=1)
doc.add_heading('4.1  MODEL ARCHITECTURE DIAGRAM', level=2)
doc.add_heading('4.2  DATA PIPELINE FLOWCHART', level=2)
doc.add_heading('4.3  USE CASE DIAGRAM', level=2)
doc.add_heading('4.4  SEQUENCE DIAGRAM', level=2)
doc.add_heading('4.5  ACTIVITY DIAGRAM', level=2)
doc.add_heading('4.6  DATA FLOW DIAGRAM', level=2)

page_break()

# ── 4.1 Model Architecture Diagram ──
doc.add_heading('4.1. Model Architecture Diagram', level=2)
add_normal(
    'A Model Architecture Diagram is a visual representation of the neural network structure, '
    'illustrating the sequence of layers, their configurations, and how data flows through the '
    'network from input to output. It is the most critical diagram for any deep learning project, '
    'as it directly communicates the design decisions behind the model.'
)
add_normal(
    'The Computer Vision Lip Reading System uses a 3D Convolutional Neural Network (3D CNN) '
    'built with TensorFlow/Keras Sequential API. The architecture is specifically designed to '
    'capture both spatial features (lip shape in each frame) and temporal features (how lip shape '
    'changes across the frame sequence). The complete model architecture is as follows:'
)
add_normal('Input Layer:', bold=True)
add_bullet('Input Shape: (22, 80, 112, 3) — 22 sequential lip-region frames, each 80×112 pixels in RGB.')
add_normal('Convolutional Block 1:', bold=True)
add_bullet('Conv3D: 8 filters, kernel size (3,3,3), activation=\'relu\', padding=\'same\', L2 regularization (0.001). '
           'Learns low-level spatiotemporal features such as lip edges, basic motion patterns.')
add_bullet('MaxPooling3D: pool size (1,2,2). Reduces spatial dimensions by half while preserving temporal resolution.')
add_normal('Convolutional Block 2:', bold=True)
add_bullet('Conv3D: 32 filters, kernel size (3,3,3), activation=\'relu\', padding=\'same\', L2 regularization (0.001). '
           'Learns mid-level features such as lip contour shapes and movement combinations.')
add_bullet('MaxPooling3D: pool size (1,2,2). Further spatial downsampling.')
add_normal('Convolutional Block 3:', bold=True)
add_bullet('Conv3D: 256 filters, kernel size (3,3,3), activation=\'relu\', padding=\'same\', L2 regularization (0.001). '
           'Captures high-level spatiotemporal patterns representing complete lip movement gestures for different words.')
add_bullet('MaxPooling3D: pool size (1,2,2). Final spatial downsampling.')
add_normal('Classification Head:', bold=True)
add_bullet('Flatten: Converts the 3D feature maps into a 1D feature vector.')
add_bullet('Dense (1024 units, activation=\'relu\'): First fully connected layer for high-dimensional feature mapping.')
add_bullet('Dropout (rate=0.5): Randomly drops 50% of neurons during training to prevent overfitting.')
add_bullet('Dense (256 units, activation=\'relu\'): Second fully connected layer for feature compression.')
add_bullet('Dropout (rate=0.5): Additional regularization.')
add_bullet('Dense (64 units, activation=\'relu\'): Third fully connected layer for final feature refinement.')
add_bullet('Dense (13 units, activation=\'softmax\'): Output layer producing probability distribution across 13 word classes.')
add_normal('Training Configuration:', bold=True)
add_bullet('Optimizer: Adam (default learning rate = 0.001)')
add_bullet('Loss Function: Categorical Cross-Entropy')
add_bullet('Epochs: 20')
add_bullet('Train/Test Split: 80% training, 20% validation')
add_normal('[INSERT Fig4.1: Model Architecture Diagram HERE]', italic=True)
add_normal('Fig4.1: Model Architecture Diagram (3D CNN)', bold=True)

page_break()

# ── 4.2 Data Pipeline Flowchart ──
doc.add_heading('4.2. Data Pipeline Flowchart', level=2)
add_normal(
    'A Data Pipeline Flowchart illustrates the end-to-end flow of data through a machine learning '
    'system — from raw data acquisition to final model inference. For ML projects, this diagram '
    'is essential as it shows the complete lifecycle of data, including collection, preprocessing, '
    'training, evaluation, and deployment stages.'
)
add_normal(
    'The Computer Vision Lip Reading System pipeline consists of five major stages:'
)
add_normal('Stage 1: Data Collection', bold=True)
add_bullet('Webcam captures raw video frames at native FPS.')
add_bullet('The facial landmark detector identifies faces and extracts 68 landmark points.')
add_bullet('Lip distance between upper and lower lip landmarks is monitored to detect speech onset/offset.')
add_bullet('When speech is detected, 22 lip-region frames (including 4 pre-speech buffer frames) are captured.')
add_bullet('Frames are saved as image files, structured numerical arrays, and video clips in labeled folders.')
add_normal('Stage 2: Data Preprocessing', bold=True)
add_bullet('Each lip frame (80×112 pixels) undergoes a standardization pipeline:')
add_bullet('Color space conversion → Contrast enhancement on the luminance channel → Back-conversion.', level=1)
add_bullet('Gaussian smoothing → Bilateral filtering → Sharpening → Final smoothing.', level=1)
add_bullet('This normalizes lip appearance across varying lighting and camera conditions.')
add_normal('Stage 3: Model Training', bold=True)
add_bullet('All preprocessed samples are loaded and organized into input tensors and one-hot encoded labels.')
add_bullet('Data is split 80/20 for training and validation with a fixed random seed for reproducibility.')
add_bullet('The 3D CNN model is compiled with the Adam optimizer and categorical cross-entropy loss.')
add_bullet('The model is trained for 20 epochs, achieving 95.7% training accuracy and 98.5% validation accuracy.')
add_bullet('Trained weights are saved for later inference use.')
add_normal('Stage 4: Model Evaluation', bold=True)
add_bullet('Confusion matrix generated to analyze per-class classification performance.')
add_bullet('ROC-AUC curves plotted for all 13 classes.')
add_bullet('Per-class precision, recall, and F1-score computed.')
add_bullet('Balanced accuracy score confirms the model performs well across all classes.')
add_normal('Stage 5: Inference', bold=True)
add_bullet('The same preprocessing pipeline used during training is applied to new input.')
add_bullet('22 preprocessed lip frames are stacked into an input tensor.')
add_bullet('The trained model outputs a probability distribution across the 13 word classes.')
add_bullet('The class with the highest probability is selected as the predicted word.')
add_bullet('The predicted word is displayed on screen (live mode) or annotated on the output video (file mode).')
add_normal('[INSERT Fig4.2: Data Pipeline Flowchart HERE]', italic=True)
add_normal('Fig4.2: Data Pipeline Flowchart', bold=True)

page_break()

# ── 4.3 Use Case Diagram ──
doc.add_heading('4.3. Use Case Diagram', level=2)
add_normal(
    'A use case diagram depicts the interactions between actors and the system, showcasing '
    'functionality from the user\'s perspective.'
)
add_normal('Actors:', bold=True)
add_bullet('Data Collector (Primary Actor): The person who records training data by speaking words into the webcam.')
add_bullet('End User (Primary Actor): The person who uses the live prediction system.')
add_bullet('Webcam Hardware (System Actor): Provides real-time video frames.')
add_bullet('Facial Landmark Detector (System Actor): Provides facial landmark coordinates.')
add_bullet('3D CNN Model (System Actor): Performs word classification.')

add_normal('Use Cases:', bold=True)
add_bullet('Calibrate Lip Distance Threshold')
add_bullet('Collect Training Data for a Word')
add_bullet('Start Live Prediction Session')
add_bullet('Detect Face in Frame')
add_bullet('Extract Lip Region')
add_bullet('Detect Speaking Activity')
add_bullet('Capture & Preprocess Lip Frames')
add_bullet('Predict Spoken Word (automated)')
add_bullet('Display Predicted Word')
add_bullet('Run File-Based Prediction')
add_bullet('Save Output Video with Annotations')
add_bullet('Reset Prediction History')
add_normal('[INSERT Fig4.3: Use Case Diagram HERE]', italic=True)
add_normal('Fig4.3: Use Case Diagram', bold=True)

page_break()

# ── 4.4 Sequence Diagram ──
doc.add_heading('4.4. Sequence Diagram', level=2)
add_normal(
    'A sequence diagram shows the sequence of messages passed between objects over time during an interaction.'
)
add_normal('Live Prediction Flow Sequence:', bold=True)
add_normal('1. User → Prediction Module: Launches the system')
add_normal('2. Prediction Module → Landmark Detector: Loads pre-trained facial landmark model')
add_normal('3. Prediction Module → Deep Learning Framework: Loads trained neural network weights')
add_normal('4. Prediction Module → Computer Vision Library: Opens webcam capture')
add_normal('5. [Loop] For each webcam frame:')
add_normal('   5.1. Webcam → Prediction Module: Returns captured frame', indent=1)
add_normal('   5.2. Prediction Module → Landmark Detector: Detects face bounding boxes', indent=1)
add_normal('   5.3. Prediction Module → Landmark Detector: Extracts 68 facial landmarks', indent=1)
add_normal('   5.4. Prediction Module: Calculates lip distance between upper and lower lip landmarks', indent=1)
add_normal('   5.5. [Decision] Lip distance exceeds speaking threshold?', indent=1)
add_normal('        Yes → Append preprocessed lip frame to current word buffer', indent=2)
add_normal('        No  → Increment silence counter', indent=2)
add_normal('   5.6. [Decision] Silence counter exceeds threshold AND total accumulated frames equal 22?', indent=1)
add_normal('        Yes → Prepend buffered pre-speech frames, form complete 22-frame input', indent=2)
add_normal('   5.7. Prediction Module → Deep Learning Framework: Classify the 22-frame input tensor', indent=1)
add_normal('   5.8. Deep Learning Framework → Prediction Module: Returns probability array for 13 classes', indent=1)
add_normal('   5.9. Prediction Module: Selects the class with highest probability as predicted word', indent=1)
add_normal('   5.10. Prediction Module → Display: Overlays predicted word on the video frame', indent=1)
add_normal('6. User exits → System terminates')
add_normal('[INSERT Fig4.4: Sequence Diagram HERE]', italic=True)
add_normal('Fig4.4: Sequence Diagram', bold=True)

page_break()

# ── 4.5 Activity Diagram ──
doc.add_heading('4.5. Activity Diagram', level=2)
add_normal(
    'An activity diagram models the flow from one activity to another within the system.'
)
add_normal('Data Collection & Prediction Activity Flow:', bold=True)
add_normal('1. [Start] → User launches the system (data collection mode or prediction mode)')
add_normal('2. → System loads facial landmark detector and trained model weights')
add_normal('3. → System opens webcam for video capture')
add_normal('4. → [Decision] Is this data collection mode?')
add_normal('     Yes → Prompt user for word label and optional lip distance threshold')
add_normal('     No (prediction mode) → Continue directly')
add_normal('5. → [Decision] Need calibration? (No custom threshold provided)')
add_normal('     Yes → Measure lip distance for 50 frames with closed mouth → Calculate average threshold')
add_normal('     No → Use custom threshold')
add_normal('6. → [Loop] For each frame from webcam:')
add_normal('     → Convert to grayscale → Detect faces → Extract landmarks')
add_normal('     → Calculate lip distance between upper and lower lip landmarks')
add_normal('     → Extract lip region from landmark coordinates → Resize to standard dimensions')
add_normal('     → Preprocess: Color space conversion → Contrast enhancement → Smoothing → Filtering → Sharpening')
add_normal('     → [Decision] Lip distance exceeds speaking threshold?')
add_normal('       Yes (Talking) → Append lip frame to current word buffer')
add_normal('       No (Not talking) → Increment silence counter')
add_normal('     → [Decision] Word complete? (silence counter exceeds threshold AND total frames equal 22)')
add_normal('       Data Collection → Save word data to the dataset directory')
add_normal('       Prediction → Feed 22 frames to model → Display predicted word')
add_normal('7. → User exits → Release webcam → [End]')
add_normal('[INSERT Fig4.5: Activity Diagram HERE]', italic=True)
add_normal('Fig4.5: Activity Diagram', bold=True)

page_break()

# ── 4.6 Data Flow Diagram ──
doc.add_heading('4.6. Data Flow Diagram', level=2)

doc.add_heading('4.6.1. Context-Level-0 DFD', level=3)
add_normal(
    'The Context-Level-0 DFD shows the entire system as a single process:'
)
add_bullet('User provides: Word label (collection mode), Start/Stop commands')
add_bullet('Lip Reading System returns to User: Predicted word, Visual annotations on webcam feed')
add_bullet('Webcam Hardware provides: Raw video frames')
add_bullet('File System exchanges: Training data files (images, numerical arrays, video clips), Model weight files')
add_normal('[INSERT Fig4.6.1: Context-Level-0 DFD HERE]', italic=True)

doc.add_heading('4.6.2. Context-Level-1 DFD', level=3)
add_normal('The Level-1 DFD breaks the system into major processes:')
add_bullet('Process 1: Face & Lip Detection – Receives raw frames, detects face using the facial landmark detector, extracts 68 landmarks, crops the lip region.')
add_bullet('Process 2: Lip Preprocessing – Applies color space conversion, contrast enhancement, smoothing, bilateral filtering, and sharpening to normalize lip frames.')
add_bullet('Process 3: Speech Activity Detection – Monitors lip distance to determine talking/not-talking state, manages frame buffers.')
add_bullet('Process 4: Frame Accumulation – Accumulates 22 preprocessed lip frames (including pre-speech buffer frames) for a single word.')
add_bullet('Process 5: Classification (Prediction Mode) – Feeds 22-frame tensor through the 3D CNN model, outputs the predicted word label.')
add_bullet('Process 6: Data Saving (Collection Mode) – Saves 22-frame sequences as structured files, images, and videos to the dataset directory.')
add_bullet('Data Store: Dataset Directory – Stores labeled word folders with frame data. Model Directory – Stores trained model weights.')
add_normal('[INSERT Fig4.6.2: Context-Level-1 DFD HERE]', italic=True)

doc.add_heading('4.6.3. Context-Level-2 DFD For Data Collection', level=3)
add_normal('Expands Process 1-4 and 6 for the collection pipeline:')
add_bullet('1.1: Webcam captures raw video frame')
add_bullet('1.2: Convert frame to grayscale for face detection')
add_bullet('1.3: Face detector identifies face bounding boxes')
add_bullet('1.4: Landmark predictor extracts 68 facial landmark coordinates')
add_bullet('2.1: Extract lip region using horizontal and vertical lip landmark coordinates')
add_bullet('2.2: Resize lip frame to standardized 112×80 pixel dimensions')
add_bullet('2.3: Apply preprocessing chain: Color space conversion → Contrast enhancement → Smoothing → Bilateral filtering → Sharpening → Final smoothing')
add_bullet('3.1: Calculate lip distance as Euclidean distance between upper and lower lip landmarks')
add_bullet('3.2: Compare with the speaking threshold')
add_bullet('4.1: Manage circular pre-speech frame buffer (capacity of 4 frames)')
add_bullet('4.2: Accumulate active speech frames during talking')
add_bullet('4.3: Prepend buffered pre-speech frames when word is complete → 22 total frames')
add_bullet('6.1: Save frame data as structured numerical array file')
add_bullet('6.2: Save individual image frames')
add_bullet('6.3: Combine frames into a video clip')
add_normal('[INSERT Fig4.6.3: Context-Level-2 DFD For Data Collection HERE]', italic=True)

doc.add_heading('4.6.4. Context-Level-2 DFD For Prediction', level=3)
add_normal('Expands Process 5 for the prediction pipeline:')
add_bullet('5.1: Reshape 22 frames into input tensor with appropriate dimensions')
add_bullet('5.2: Feed input tensor through the trained 3D CNN model → Obtain probability distribution across 13 classes')
add_bullet('5.3: Select the class with the highest probability as the predicted class index')
add_bullet('5.4: Check prediction history → Skip if the same word was already predicted')
add_bullet('5.5: Map predicted class index to its corresponding word label')
add_bullet('5.6: Display the predicted word as an overlay on the video frame')
add_normal('[INSERT Fig4.6.4: Context-Level-2 DFD For Prediction HERE]', italic=True)

page_break()

# ════════════════════════════════════════════════════════════════════════
#  CHAPTER 5: DATA DICTIONARY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('5    DATA DICTIONARY', level=1)
doc.add_heading('5.1  OVERVIEW', level=2)
doc.add_heading('5.2  DATA SCHEMA', level=2)
doc.add_heading('5.3  DATA INTEGRITY AND PROCESSING', level=2)
doc.add_heading('5.4  DATASET DESCRIPTION AND STATISTICS', level=2)

page_break()

# ── 5.1 Overview ──
doc.add_heading('5.1. Overview', level=2)
add_normal(
    'The Computer Vision Lip Reading System stores and processes data in two forms: (1) the training '
    'dataset stored as files on disk in a structured dataset directory, and (2) the in-memory data '
    'structures used during real-time collection and prediction. Unlike database-driven applications, '
    'this system uses the file system as its primary storage mechanism for training data, and numerical '
    'arrays as the in-memory data format for model input/output. The model weights are stored in '
    'HDF5 format, and the facial landmark model is stored as a pre-trained binary file.'
)

page_break()

# ── 5.2 Data Schema ──
doc.add_heading('5.2. Data Schema', level=2)

doc.add_heading('5.2.1. Collected Data Structure', level=3)
add_normal(
    'Each word sample is stored in a dedicated folder within the dataset directory, named by '
    'word label and sample number (e.g., hello_1/, cat_23/, bye_5/).'
)
add_table(
    ['File/Item', 'Data Type', 'Description'],
    [
        ['Folder Name', 'String (e.g., "hello_42")', 'Named as {label}_{index}. Contains all data for one spoken word sample.'],
        ['Frame Data File', 'JSON (3D array)', 'Serialized 3D numerical array: [22 frames][80 rows][112 cols][3 channels]. Each value is an integer 0-255 (RGB pixel value).'],
        ['Frame Images (0 to 21)', 'PNG images (80×112)', '22 individual lip-region frames saved as image files. Frame 0 is the earliest frame, frame 21 is the latest.'],
        ['Video Clip', 'MP4 video', 'All 22 frames combined into a short video clip at the webcam\'s native FPS for visual review.'],
    ]
)
doc.add_paragraph()
add_normal('Total Dataset Statistics:', bold=True)
add_table(
    ['Metric', 'Value'],
    [
        ['Total Samples', '~700 video clips'],
        ['Total Data Size', '~3 GB'],
        ['Word Classes', '13'],
        ['Frames per Sample', '22'],
        ['Frame Dimensions', '80 (height) × 112 (width) × 3 (RGB channels)'],
        ['Words', '"here", "is", "a", "demo", "can", "you", "read", "my", "lips", "cat", "dog", "hello", "bye"'],
    ]
)

doc.add_heading('5.2.2. Model Input Tensor', level=3)
add_normal(
    'The 3D CNN model expects input in the following tensor format:'
)
add_table(
    ['Dimension', 'Size', 'Description'],
    [
        ['Batch', '1 (inference) / 16 (training)', 'Number of samples per batch'],
        ['Frames (Temporal)', '22', 'Number of sequential lip-region frames per word'],
        ['Height (Spatial)', '80', 'Height of each preprocessed lip frame in pixels'],
        ['Width (Spatial)', '112', 'Width of each preprocessed lip frame in pixels'],
        ['Channels', '3', 'RGB color channels'],
    ]
)
add_normal('Full input shape: (batch_size, 22, 80, 112, 3)', bold=True)

doc.add_paragraph()
add_normal('Label Encoding Mapping:', bold=True)
add_table(
    ['Encoded Value', 'Word Label'],
    [
        ['0', 'a'], ['1', 'bye'], ['2', 'can'], ['3', 'cat'],
        ['4', 'demo'], ['5', 'dog'], ['6', 'hello'], ['7', 'here'],
        ['8', 'is'], ['9', 'lips'], ['10', 'my'], ['11', 'read'], ['12', 'you'],
    ]
)

page_break()

# ── 5.3 Data Integrity and Processing ──
doc.add_heading('5.3. Data Integrity and Processing', level=2)
add_normal(
    'Frame Count Consistency: Every word sample is guaranteed to contain exactly 22 frames. '
    'The data collection module enforces this through a combination of pre-speech buffer frames (4 frames '
    'of context), active speech frames, and silence detection logic. Samples that do not '
    'meet the exact 22-frame requirement are discarded during collection.'
)
add_normal(
    'Spatial Consistency: All lip frames are resized to exactly 112×80 pixels using standard image resizing, '
    'regardless of the original lip region size detected by the facial landmarks. This ensures '
    'uniform tensor dimensions for the 3D CNN input.'
)
add_normal(
    'Preprocessing Pipeline: Every frame undergoes the same preprocessing chain:'
)
add_bullet('LAB Color Space Conversion for luminance-chrominance separation')
add_bullet('CLAHE Contrast Enhancement (clip limit of 3.0, tile grid of 3×3) on the luminance channel')
add_bullet('Conversion back to the original color space')
add_bullet('Gaussian Blur (7×7 kernel) for initial noise reduction')
add_bullet('Bilateral Filtering for edge-preserving smoothing')
add_bullet('Sharpening via a custom convolution kernel')
add_bullet('Final Gaussian Blur (5×5 kernel) for subtle smoothing')
add_normal(
    'Data Types: Raw frames are stored as lists of integers (0-255) in serialized format. For model '
    'input, they are converted to floating-point numerical arrays. The model uses 32-bit floating point '
    'internally for all computations.'
)
add_normal(
    'Train/Test Split: The dataset is split 80% training / 20% validation using a stratified '
    'random split with a fixed random seed for reproducibility.'
)

page_break()

# ── 5.4 Dataset Description and Statistics ──
doc.add_heading('5.4. Dataset Description and Statistics', level=2)
add_normal(
    'This section provides a comprehensive description of the dataset used for training and evaluating '
    'the lip reading model. Since no suitable publicly available dataset existed for this specific '
    'word-level lip reading task with the required preprocessing format, a custom dataset was created '
    'from scratch using the data collection pipeline described in earlier sections.'
)

doc.add_heading('5.4.1. Dataset Overview', level=3)
add_normal(
    'The dataset consists of approximately 700 video clips of individual words being spoken, captured '
    'using a standard webcam under controlled indoor lighting conditions. The dataset was manually '
    'collected and labeled, with each sample representing a single spoken word captured as a sequence '
    'of 22 preprocessed lip-region frames. The total dataset size is approximately 3 GB, stored across '
    'individually labeled folders in the project\'s dataset directory. The dataset is also publicly '
    'available on Kaggle for reproducibility and further research.'
)
add_normal(
    'The dataset comprises 13 distinct word classes, carefully selected to include a mix of common '
    'English words and demonstrative phrases. The vocabulary was chosen to enable the construction '
    'of simple sentences (e.g., "here is a demo", "can you read my lips") while also including '
    'visually distinct words (e.g., "cat", "dog", "hello", "bye") for robust classification evaluation.'
)

doc.add_heading('5.4.2. Per-Class Sample Distribution', level=3)
add_normal(
    'The following table presents the number of collected samples for each word class in the dataset. '
    'The distribution is approximately balanced, with each class containing between 49 and 61 samples, '
    'which helps reduce class imbalance bias during model training.'
)
add_table(
    ['Word Class', 'Number of Samples', 'Percentage of Dataset'],
    [
        ['a', '53', '7.8%'],
        ['bye', '52', '7.6%'],
        ['can', '51', '7.5%'],
        ['cat', '50', '7.3%'],
        ['demo', '52', '7.6%'],
        ['dog', '50', '7.3%'],
        ['hello', '49', '7.2%'],
        ['here', '61', '8.9%'],
        ['is', '54', '7.9%'],
        ['lips', '53', '7.8%'],
        ['my', '53', '7.8%'],
        ['read', '54', '7.9%'],
        ['you', '51', '7.5%'],
        ['Total', '683', '100%'],
    ]
)
add_normal(
    'The approximate balance across classes (standard deviation of \u22483.2 samples) indicates that the '
    'dataset is well-distributed, minimizing the risk of the model developing a bias toward any '
    'particular word class. The "here" class has the highest sample count (61), while "hello" has '
    'the lowest (49), representing only a 12-sample difference across the extremes.'
)
add_normal('[INSERT Fig: Class Distribution Histogram HERE]', italic=True)

doc.add_heading('5.4.3. Dataset Characteristics', level=3)
add_normal(
    'The following table summarizes the key characteristics of the dataset:'
)
add_table(
    ['Characteristic', 'Value'],
    [
        ['Dataset Name', 'Custom Lip Reading Dataset'],
        ['Public Availability', 'Available on Kaggle (best-lip-reading-dataset)'],
        ['Total Samples', '683 video clips'],
        ['Number of Classes', '13 word classes'],
        ['Total Dataset Size', '~3 GB'],
        ['Frames per Sample', '22 (fixed)'],
        ['Frame Resolution', '80 × 112 pixels (H × W)'],
        ['Color Space', 'RGB (3 channels)'],
        ['Input Tensor Shape', '(22, 80, 112, 3) per sample'],
        ['Collection Method', 'Webcam-based with automated speech detection'],
        ['Labeling Method', 'Manual labeling at time of collection'],
        ['Number of Speakers', '2 (the project creators)'],
        ['Recording Environment', 'Controlled indoor lighting'],
        ['Train/Validation Split', '80% / 20% (stratified, fixed random seed)'],
        ['Training Set Size', '~546 samples'],
        ['Validation Set Size', '~137 samples'],
    ]
)

doc.add_heading('5.4.4. Data Collection Methodology', level=3)
add_normal(
    'The data collection process was designed to capture natural lip movements during speech while '
    'maintaining consistency across all samples. The methodology involved the following steps:'
)
add_normal('1. Subject Positioning:', bold=True)
add_normal(
    'The speaker sits approximately 50-70 cm from the webcam, facing directly toward the camera '
    'with adequate frontal lighting to ensure clear facial visibility. The environment is kept '
    'consistent across all recording sessions to minimize variance in background and lighting.'
)
add_normal('2. Lip Distance Calibration:', bold=True)
add_normal(
    'Before each recording session, the system performs a calibration step. The speaker keeps '
    'their mouth closed while the system measures the resting lip distance (the Euclidean distance '
    'between the upper and lower lip landmarks) over 50 frames to establish a personalized '
    'speaking threshold. This threshold is used to detect when the speaker begins and stops talking.'
)
add_normal('3. Word Recording:', bold=True)
add_normal(
    'The speaker is prompted to speak a specific word. The system monitors the lip distance in '
    'real-time and begins recording when the distance exceeds the calibrated threshold (indicating '
    'the speaker has opened their mouth). Recording continues until the speaker closes their mouth '
    'for a sustained period (10 consecutive frames below threshold), at which point the system '
    'determines the word has been completed.'
)
add_normal('4. Frame Assembly:', bold=True)
add_normal(
    'Each captured word is assembled into a fixed 22-frame sequence. This includes 4 pre-speech '
    'buffer frames (captured before the speaking threshold was crossed) and the active speech '
    'frames. If additional frames are needed to reach the 22-frame target, the system continues '
    'capturing until the requirement is met. This ensures temporal context is preserved from '
    'before the onset of speech.'
)
add_normal('5. Preprocessing and Storage:', bold=True)
add_normal(
    'Each of the 22 lip-region frames undergoes the standardized preprocessing pipeline (color space '
    'conversion, contrast enhancement, smoothing, and sharpening) before being saved. The processed '
    'frames are stored in three formats: individual image files for visual inspection, a serialized '
    'numerical array for model training, and a video clip for human review of the captured sequence.'
)
add_normal('6. Quality Control:', bold=True)
add_normal(
    'Each recorded sample is visually verified through the saved video clip. Samples with poor '
    'face detection, excessive motion blur, or incomplete word articulation are manually discarded '
    'and re-recorded to maintain dataset quality.'
)

page_break()

# ════════════════════════════════════════════════════════════════════════
#  CHAPTER 6: RESULT, DISCUSSION AND CONCLUSION
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('6    RESULT, DISCUSSION AND CONCLUSION', level=1)
doc.add_heading('6.1  RESULT', level=2)
doc.add_heading('6.2  CONCLUSION', level=2)

page_break()

# ── 6.1 Result ──
doc.add_heading('6.1. Result', level=2)
add_normal(
    'The Computer Vision Lip Reading System has been successfully developed and tested, '
    'demonstrating the following functional outcomes:'
)
add_normal('1. Custom Dataset Creation:', bold=True)
add_normal(
    'The data collection pipeline successfully captured approximately 700 video clips '
    'across 13 word classes, totaling ~3 GB of data. Each clip contains exactly 22 preprocessed '
    'lip-region frames at 80×112 pixels. The automated speech detection based on lip-distance '
    'thresholds reliably distinguished speaking from non-speaking states.'
)
add_normal('2. Model Training Performance:', bold=True)
add_normal(
    'The 3D CNN model (3 convolutional layers with L2 regularization, 3 fully connected layers with dropout) was '
    'trained for 20 epochs with Adam optimizer and categorical cross-entropy loss, achieving:'
)
add_bullet('Training Accuracy: 95.7%')
add_bullet('Validation Accuracy: 98.5%')
add_bullet('The validation accuracy exceeding training accuracy suggests effective regularization '
           'and a well-generalized model that is not overfitting.')

add_normal('3. Classification Metrics:', bold=True)
add_normal(
    'Per-class precision, recall, and F1-score were computed using standard evaluation metrics. '
    'The confusion matrix heatmap shows strong diagonal dominance, indicating that the model correctly '
    'classifies the vast majority of samples for each of the 13 word classes. The ROC-AUC curve '
    'demonstrates near-perfect area under the curve for all classes.'
)
add_normal('[INSERT Training Accuracy/Loss Graph, Confusion Matrix, ROC-AUC Curve, Metrics Table HERE]', italic=True)

add_normal('4. Real-Time Live Prediction:', bold=True)
add_normal(
    'The live prediction module successfully performs real-time word prediction from a webcam feed. '
    'The system detects when the user speaks, captures 22 lip frames, feeds them through the model, '
    'and displays the predicted word on screen. A duplicate prevention mechanism ensures that the same word '
    'is not predicted consecutively, and the user can reset the prediction history for a new sentence.'
)

add_normal('5. File-Based Prediction:', bold=True)
add_normal(
    'The file-based prediction module successfully processes pre-recorded video files, performs frame-by-frame '
    'lip reading, and saves annotated output videos with predicted word overlays. This enables batch '
    'processing and demonstration recording.'
)
add_normal('[INSERT Screenshots – Data Collection Mode, Live Prediction Mode, Sample Lip Frames HERE]', italic=True)

page_break()

# ── 6.2 Conclusion ──
doc.add_heading('6.2. Conclusion', level=2)
add_normal(
    'This project successfully demonstrates that visual-only lip reading using 3D Convolutional Neural '
    'Networks is a viable approach for word-level speech recognition. By combining facial landmark '
    'detection, robust image preprocessing, and a well-designed 3D CNN architecture, the system achieves '
    '98.5% validation accuracy across 13 word classes—without using any audio information.'
)
add_normal('The project validates several key technical achievements:')
add_bullet('A custom-built data collection pipeline can efficiently generate high-quality lip reading '
           'training data using only a standard webcam and a pre-trained facial landmark detector.')
add_bullet('3D CNNs are effective at capturing spatiotemporal features from lip movement sequences, '
           'enabling accurate word classification from visual data alone.')
add_bullet('The multi-stage preprocessing pipeline (color space conversion, contrast enhancement, smoothing, '
           'bilateral filtering, and sharpening) effectively normalizes lip frame appearance across varying lighting conditions.')
add_bullet('The complete pipeline from data collection → training → real-time inference can run on '
           'standard consumer hardware without requiring specialized equipment.')

add_normal('Future enhancements planned for subsequent iterations include:', bold=True)
add_bullet('Expanding the vocabulary beyond 13 words to support a larger set of common English words.')
add_bullet('Implementing sentence-level lip reading by combining the word-level model with a language model '
           'for context-aware prediction.')
add_bullet('Adding speaker-independent training by collecting data from multiple speakers to improve '
           'generalization across different face shapes and speaking styles.')
add_bullet('Implementing a graphical user interface (GUI) for easier usage by non-technical users.')
add_bullet('Exploring transfer learning from larger pre-trained video models (e.g., I3D, SlowFast) to '
           'improve accuracy with less training data.')
add_bullet('Deploying as a web application or mobile app for broader accessibility.')
add_bullet('Integrating with assistive technology frameworks for hearing-impaired users.')

add_normal(
    'By bridging the gap between audio-dependent speech recognition and visual-only lip reading, this '
    'project contributes to the growing body of work demonstrating that computer vision and deep learning '
    'are powerful tools for building accessible communication technology.'
)

page_break()

# ════════════════════════════════════════════════════════════════════════
#  CHAPTER 7: BIBLIOGRAPHY
# ════════════════════════════════════════════════════════════════════════
doc.add_heading('7    BIBLIOGRAPHY', level=1)
doc.add_heading('7.1  REFERENCES', level=2)

page_break()

doc.add_heading('7.1. References', level=2)

doc.add_heading('7.1.1. Academic Papers & Journals', level=3)
add_normal(
    '[LipNet] Y. M. Assael, B. Shillingford, S. Whiteson, and N. de Freitas, "LipNet: End-to-End '
    'Sentence-level Lipreading," arXiv preprint arXiv:1611.01599, 2016.'
)
add_normal(
    '[Lip Reading in the Wild] J. S. Chung, A. Senior, O. Vinyals, and A. Zisserman, "Lip Reading '
    'Sentences in the Wild," Proceedings of the IEEE Conference on Computer Vision and Pattern '
    'Recognition (CVPR), pp. 6447-6456, 2017.'
)
add_normal(
    '[C3D] D. Tran, L. Bourdev, R. Fergus, L. Torresani, and M. Paluri, "Learning Spatiotemporal '
    'Features with 3D Convolutional Networks," Proceedings of the IEEE International Conference on '
    'Computer Vision (ICCV), pp. 4489-4497, 2015.'
)
add_normal(
    '[Facial Landmarks] V. Kazemi and J. Sullivan, "One Millisecond Face Alignment with an Ensemble '
    'of Regression Trees," Proceedings of the IEEE Conference on Computer Vision and Pattern '
    'Recognition (CVPR), pp. 1867-1874, 2014.'
)
add_normal(
    '[Audio-Visual SR] G. Potamianos, C. Neti, G. Gravier, A. Garg, and A. W. Senior, "Recent '
    'Advances in the Automatic Recognition of Audiovisual Speech," Proceedings of the IEEE, vol. 91, '
    'no. 9, pp. 1306-1326, 2003.'
)
add_normal(
    '[McGurk Effect] H. McGurk and J. MacDonald, "Hearing lips and seeing voices," Nature, vol. 264, '
    'no. 5588, pp. 746-748, 1976.'
)
add_normal(
    '[GRID Corpus] M. Cooke, J. Barker, S. Cunningham, and X. Shao, "An audio-visual corpus for '
    'speech perception and automatic speech recognition," Journal of the Acoustical Society of America, '
    'vol. 120, no. 5, pp. 2421-2424, 2006.'
)
add_normal(
    '[AlexNet] A. Krizhevsky, I. Sutskever, and G. E. Hinton, "ImageNet Classification with Deep '
    'Convolutional Neural Networks," Advances in Neural Information Processing Systems (NIPS), '
    'pp. 1097-1105, 2012.'
)
add_normal(
    '[WHO Hearing Loss] World Health Organization, "World Report on Hearing," Geneva: WHO, 2021. '
    '(Evidence that 430 million people globally have disabling hearing loss).'
)

doc.add_heading('7.1.2. Technical Documentation', level=3)
add_normal(
    'TensorFlow Documentation. [Online]. Available: https://www.tensorflow.org/api_docs. '
    '[Accessed: Jan 2025]. (Deep learning framework for model building, training, and inference).'
)
add_normal(
    'Keras Documentation. [Online]. Available: https://keras.io/api/. '
    '[Accessed: Jan 2025]. (High-level neural network API: Sequential, Conv3D, Dense, Dropout).'
)
add_normal(
    'OpenCV Documentation. [Online]. Available: https://docs.opencv.org/4.x/. '
    '[Accessed: Jan 2025]. (Computer vision library: VideoCapture, image processing, color space conversion).'
)
add_normal(
    'dlib Documentation. [Online]. Available: http://dlib.net/python/index.html. '
    '[Accessed: Jan 2025]. (Face detection and 68-point facial landmark prediction).'
)
add_normal(
    'NumPy Documentation. [Online]. Available: https://numpy.org/doc/stable/. '
    '[Accessed: Jan 2025]. (Numerical computing library for array operations).'
)
add_normal(
    'scikit-learn Documentation. [Online]. Available: https://scikit-learn.org/stable/. '
    '[Accessed: Jan 2025]. (Machine learning utilities: train_test_split, classification metrics).'
)

doc.add_heading('7.1.3. Web Resources & Libraries', level=3)
add_normal(
    'dlib – Face Landmark Detection. [Online]. Available: http://dlib.net/. '
    '(C++ toolkit with Python bindings for face detection and shape prediction).'
)
add_normal(
    'dlib Shape Predictor 68 Landmarks. [Online]. Available: '
    'https://github.com/davisking/dlib-models. (Pre-trained facial landmark model used in this project).'
)
add_normal(
    'OpenCV – Open Source Computer Vision Library. [Online]. Available: https://opencv.org/. '
    '(Industry-standard computer vision library used for image/video processing).'
)
add_normal(
    'Kaggle Lip Reading Dataset. [Online]. Available: '
    'https://www.kaggle.com/datasets/allenye66/best-lip-reading-dataset. '
    '(The dataset created for this project, hosted on Kaggle).'
)
add_normal(
    'imageio – Python Library for Image/Video I/O. [Online]. Available: https://imageio.readthedocs.io/. '
    '(Used for reading images and writing MP4 videos from frame sequences).'
)

doc.add_heading('7.1.4. Similar Systems (Case Studies)', level=3)
add_normal(
    'LipNet – End-to-End Lip Reading. [Online]. Available: https://github.com/rizkiarm/LipNet. '
    '(Open-source implementation of the LipNet sentence-level lip reading model).'
)
add_normal(
    'Google Speech-to-Text. [Online]. Available: https://cloud.google.com/speech-to-text. '
    '(Audio-based speech recognition API — does not support visual lip reading).'
)
add_normal(
    'Liopa – Visual Speech Recognition. [Online]. Available: https://www.liopa.ai/. '
    '(Commercial lip reading technology for healthcare, developed in the UK).'
)
add_normal(
    'Mozilla DeepSpeech. [Online]. Available: https://github.com/mozilla/DeepSpeech. '
    '(Open-source audio speech-to-text engine — audio-only, no visual component).'
)
add_normal(
    'Computer-Vision-Lip-Reading (v1.0). [Online]. Available: '
    'https://github.com/allenye66/Computer-Vision-Lip-Reading. '
    '(The predecessor project that this system builds upon, using a different strategy).'
)

# ════════════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════════════
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    'docus',
    'LipReading_Project_Report.docx'
)
doc.save(output_path)
print(f"Report saved to: {output_path}")